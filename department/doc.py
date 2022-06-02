from docx import Document
from docx.shared import Inches


def create_doc(id, employee, ITN, insPolicy, snils,
               phoneNum, bornDate, sex,
               fact_living_place, family,
               passport, education, contract, empBook):

    document = Document()
    document.add_heading(f'ЛИЧНОЕ ДЕЛО № {id if id else "-"}', 0)
    document.add_paragraph(
        f'1. Имя: {employee.firstName if employee.firstName else "-"}')
    document.add_paragraph(
        f'2. Фамилия: {employee.middleName if employee.middleName else "-"}')
    document.add_paragraph(
        f'3. Отчество: {employee.lastName if employee.lastName else "-"}')
    document.add_paragraph(f'4. Пол: {sex if sex else "-"}')
    document.add_paragraph(
        f'5. Номер телефона: {phoneNum if phoneNum else "-"}')
    document.add_paragraph(
        f'6. Место проживания: {fact_living_place if fact_living_place else "-"}')
    document.add_paragraph(
        f'7. Дата рождения: {bornDate.strftime("%d-%m-%Y") if bornDate else "-"}')
    document.add_paragraph(
        f'8. Место проживания: {fact_living_place if fact_living_place else "-"}')
    document.add_paragraph(f'9. ИНН: {ITN if ITN else "-"}')
    document.add_paragraph(f'10. СНИЛС {snils if snils else "-"}')
    document.add_paragraph(
        f'11. Номер мед. полиса: {insPolicy if insPolicy else "-"}')
    document.add_paragraph(f'12. Семья: {"" if family else "-"}')

    if family:
        records_family = (
            (family.marStatus if family.marStatus else '-',
             family.maidenName if family.maidenName else '-'),
        )
        table_family = document.add_table(rows=1, cols=2)
        hdr_cells_family = table_family.rows[0].cells
        hdr_cells_family[0].text = 'Семейное положение'
        hdr_cells_family[1].text = 'Девичья фамилия'
        for r1, r2 in records_family:
            row_cells_family = table_family.add_row().cells
            row_cells_family[0].text = str(r1) if r1 else '-'
            row_cells_family[1].text = str(r2) if r2 else '-'

    document.add_paragraph(f'14. Паспорт: {"" if passport else "-"}')

    if passport:
        records_passport = (
            (passport.series if passport.series else '-',
             passport.number if passport.number else '-',
             passport.code if passport.code else '-',
             passport.place if passport.place else '-',
             passport.date.strftime("%d-%m-%Y") if passport.date else '-',
             passport.bornPlace if passport.bornPlace else '-',
             passport.living_place if passport.living_place else '-'),
        )
        table_passport = document.add_table(rows=1, cols=7)
        hdr_cells_passport = table_passport.rows[0].cells
        hdr_cells_passport[0].text = 'Серия'
        hdr_cells_passport[1].text = 'Номер'
        hdr_cells_passport[2].text = 'Код подразделения'
        hdr_cells_passport[3].text = 'Кем выдан'
        hdr_cells_passport[4].text = 'Дата выдачи'
        hdr_cells_passport[5].text = 'Место рождения'
        hdr_cells_passport[6].text = 'Место прописки'
        for r1, r2, r3, r4, r5, r6, r7 in records_passport:
            row_cells_passport = table_passport.add_row().cells
            row_cells_passport[0].text = str(r1) if r1 else '-'
            row_cells_passport[1].text = str(r2) if r2 else '-'
            row_cells_passport[2].text = str(r3) if r3 else '-'
            row_cells_passport[3].text = str(r4) if r4 else '-'
            row_cells_passport[4].text = str(r5) if r5 else '-'
            row_cells_passport[5].text = str(r6) if r6 else '-'
            row_cells_passport[6].text = str(r7) if r7 else '-'

    document.add_paragraph(f'15. Образование: {"" if education else "-"}')

    if education:
        table_education = document.add_table(rows=1, cols=7)
        hdr_cells_education = table_education.rows[0].cells
        hdr_cells_education[0].text = 'Наименование образовательного учреждения'
        hdr_cells_education[1].text = 'Наименование документа об образовании'
        hdr_cells_education[2].text = 'Серия'
        hdr_cells_education[3].text = 'Номер'
        hdr_cells_education[4].text = 'Направление или специальность'
        hdr_cells_education[5].text = 'Квалификация'
        hdr_cells_education[6].text = 'Год окончания'
        for el in education.all():
            records_education = (
                (el.nameIns if el.nameIns else '-',
                 el.nameDoc if el.nameDoc else '-',
                 el.series if el.series else '-',
                 el.number if el.number else '-',
                 el.speciality if el.speciality else '-',
                 el.qalification if el.qalification else '-',
                 el.graduationYear if el.graduationYear else '-'
                 ),
            )

            for r1, r2, r3, r4, r5, r6, r7 in records_education:
                row_cells_education = table_education.add_row().cells
                row_cells_education[0].text = str(r1) if r1 else '-'
                row_cells_education[1].text = str(r2) if r2 else '-'
                row_cells_education[2].text = str(r3) if r3 else '-'
                row_cells_education[3].text = str(r4) if r4 else '-'
                row_cells_education[4].text = str(r5) if r5 else '-'
                row_cells_education[5].text = str(r6) if r6 else '-'
                row_cells_education[6].text = str(r7) if r7 else '-'

    document.add_paragraph(f'15. Трудовой договор: {"" if contract else "-"}')

    if contract:
        records_contract = (
            (contract.payment if contract.payment else '-',
             contract.type if contract.type else '-',
             contract.start_date.strftime(
                 "%d-%m-%Y") if contract.start_date else '-',
             contract.end_date.strftime(
                 "%d-%m-%Y") if contract.end_date else '-'
             ),
        )
        table_contract = document.add_table(rows=1, cols=4)
        hdr_cells_contract = table_contract.rows[0].cells
        hdr_cells_contract[0].text = 'Зарплата'
        hdr_cells_contract[1].text = 'Тип'
        hdr_cells_contract[2].text = 'Дата начала действия договора'
        hdr_cells_contract[3].text = 'Дата конца действия договора'
        for r1, r2, r3, r4 in records_contract:
            row_cells_contract = table_contract.add_row().cells
            row_cells_contract[0].text = str(r1) if r1 else '-'
            row_cells_contract[1].text = str(r2) if r2 else '-'
            row_cells_contract[2].text = str(r3) if r3 else '-'
            row_cells_contract[3].text = str(r4) if r4 else '-'

    document.add_paragraph(f'16. Трудовая книжка: {"" if empBook else "-"}')

    if empBook:
        records_empBook = (
            (empBook.give_date.strftime("%d-%m-%Y") if empBook.give_date else '-',
             empBook.position if empBook.position else '-',
             empBook.work_start_date.strftime(
                 "%d-%m-%Y") if empBook.work_start_date else '-',
             empBook.work_end_date.strftime(
                 "%d-%m-%Y") if empBook.work_end_date else '-',
             empBook.work_experience if empBook.work_experience else '-'
             ),
        )
        table_empBook = document.add_table(rows=1, cols=5)
        hdr_cells_empBook = table_empBook.rows[0].cells
        hdr_cells_empBook[0].text = 'Дата выдачи'
        hdr_cells_empBook[1].text = 'Должность'
        hdr_cells_empBook[2].text = 'Дата принятия на работу'
        hdr_cells_empBook[3].text = 'Дата увольнения или выхода на пенсию'
        hdr_cells_empBook[4].text = 'Стаж работы'
        for r1, r2, r3, r4, r5 in records_empBook:
            row_cells_empBook = table_empBook.add_row().cells
            row_cells_empBook[0].text = str(r1) if r1 else '-'
            row_cells_empBook[1].text = str(r2) if r2 else '-'
            row_cells_empBook[2].text = str(r3) if r3 else '-'
            row_cells_empBook[3].text = str(r4) if r4 else '-'
            row_cells_empBook[4].text = str(r5) if r5 else '-'

    # document.
    # print(f"{id=}", f"{employee=} {employee.firstName=}", f"{ITN=}", f"{insPolicy=}", f"{snils=}",
    #       phoneNum, f"{bornDate=}", f"{sex=}",
    #       fact_living_place, f"{family=}",
    #       passport, f"{education=}", f"{contract=}", f"{empBook=}")
    # id=1 employee=<Employee: user Наталья Васильева Юрьевна>
    # ITN=498390608771 insPolicy=10000000000000123 snils='123-123-123 01' phoneNum='+79321762312'
    # bornDate=datetime.date(1966, 1, 26) sex='Ж' fact_living_place='Россия, г. Нижневартовск, Зеленая ул., д. 5 кв.23'
    # family=None 4445 219261
    # education=<django.db.models.fields.related_descriptors.create_forward_many_to_many_manager.<locals>.ManyRelatedManager object at 0x109ba75e0>
    # contract=<Contract: ID: 1> empBook=<EmpBook: ID: 2>

    # p = document.add_paragraph('1. A plain paragraph having some ')
    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True

    # document.add_heading('Heading, level 1', level=1)
    # document.add_paragraph('Intense quote', style='Intense Quote')

    # document.add_paragraph(
    #     'first item in unordered list', style='List Bullet'
    # )
    # document.add_paragraph(
    #     'first item in ordered list', style='List Number'
    # )

    # records = (
    #     (3, '101', 'Spam'),
    #     (7, '422', 'Eggs'),
    #     (4, '631', 'Spam, spam, eggs, and spam')
    # )

    # table = document.add_table(rows=1, cols=3)
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Qty'
    # hdr_cells[1].text = 'Id'
    # hdr_cells[2].text = 'Desc'
    # for qty, id, desc in records:
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = str(qty)
    #     row_cells[1].text = id
    #     row_cells[2].text = desc

    # document.add_page_break()

    return document  # .save('demo.docx')
